import gradio as gr
from sentence_transformers import SentenceTransformer, util
import docx
import os
from PyPDF2 import PdfReader
import re
from datetime import datetime

# Load pre-trained model for sentence embedding
model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

# Define maximum number of resumes
MAX_RESUMES = 10

# Function to load job description from file path
def load_job_description(job_desc_file):
    if not os.path.exists(job_desc_file):
        return "Job description file not found."
    with open(job_desc_file, 'r') as file:
        job_description = file.read()
    if not job_description.strip():
        return "Job description is empty."
    return job_description

# Function to load offer letter template
def load_offer_letter_template(template_file):
    return docx.Document(template_file)
# Function to check similarity between resumes and job description
def check_similarity(job_description, resume_files):
    results = []
    job_emb = model.encode(job_description, convert_to_tensor=True)

    for resume_file in resume_files:
        resume_text = extract_text_from_resume(resume_file)
        if not resume_text:
            results.append((resume_file.name, 0, "Not Eligible", None))
            continue
        resume_emb = model.encode(resume_text, convert_to_tensor=True)
        similarity_score = util.pytorch_cos_sim(job_emb, resume_emb)[0][0].item()

        # Set a higher similarity threshold for eligibility
        if similarity_score >= 0.50:
            candidate_name = extract_candidate_name(resume_text)
            results.append((resume_file.name, similarity_score, "Eligible", candidate_name))
        else:
            results.append((resume_file.name, similarity_score, "Not Eligible", None))

    return results
 
# Extract text from resume (handles .txt, .pdf, .docx)
def extract_text_from_resume(resume_file):
    file_extension = os.path.splitext(resume_file)[1].lower()
    if file_extension not in ['.txt', '.pdf', '.docx']:
        return "Unsupported file format"
    if file_extension == '.txt':
        return read_text_file(resume_file)
    elif file_extension == '.pdf':
        return read_pdf_file(resume_file)
    elif file_extension == '.docx':
        return read_docx_file(resume_file)
    return "Failed to read the resume text."
    
    # Corrected indentation
def my_function(file_path):
    with open(file_path, 'r') as file:  # Properly indented
        data = file.read()
        return data

def read_pdf_file(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def read_docx_file(file_path):
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text
    # Extract candidate name from resume text
def extract_candidate_name(resume_text):
    name_pattern = re.compile(r'\b([A-Z][a-z]+ [A-Z][a-z]+)\b')
    matches = name_pattern.findall(resume_text)
    if matches:
        return matches[0]  # Returns the first match
    return "Unknown Candidate"
    # Create an offer letter
def create_offer_letter(candidate_name, job_title, company_name, joining_date, template_doc):
    new_doc = docx.Document()

    # Copy the template paragraphs to the new document
    for paragraph in template_doc.paragraphs:
        new_paragraph = new_doc.add_paragraph(paragraph.text)

        # Replace placeholders with actual values
        if '{{ candidate_name }}' in new_paragraph.text:
            new_paragraph.text = new_paragraph.text.replace('{{ candidate_name }}', candidate_name)
        if '{{ role }}' in new_paragraph.text:
            new_paragraph.text = new_paragraph.text.replace('{{ role }}', job_title)
        if '{{ joining_date }}' in new_paragraph.text:
            new_paragraph.text = new_paragraph.text.replace('{{ joining_date }}', joining_date)

    # Save the new offer letter with a unique name
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    output_filename = f"/tmp/{candidate_name}_{timestamp}_Offer_Letter.docx"
    new_doc.save(output_filename)
    return output_filename
    # Schedule interview with AM/PM format
def schedule_interview(candidate_name, interview_date, interview_time):
    try:
        interview_datetime = datetime.strptime(f"{interview_date} {interview_time}", '%Y-%m-%d %I:%M %p')
        message = f"Interview scheduled for {candidate_name} on {interview_datetime.strftime('%Y-%m-%d at %I:%M %p')}"
        return message
    except Exception as e:
        return f"Error in scheduling interview: {str(e)}"
        # Validate date and time with AM/PM
def validate_date_time(date_str, time_str):
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        time_obj = datetime.strptime(time_str, "%I:%M %p")
        return True, date_obj, time_obj
    except ValueError:
        return False, None, None
        # Main processing function
def process_files(job_desc, template, resumes, interview_choice, interview_date, interview_time, generate_offer_choice, role, joining_date, candidate_name):
    try:
        # Check if the number of resumes is within the allowed limit
        if len(resumes) > MAX_RESUMES:
            return "Please upload no more than 10 resumes."

        # Check if all necessary files are provided
        if not job_desc or not template or not resumes:
            return "Please provide all necessary files."

        # Load the job description and offer letter template
        job_desc_text = load_job_description(job_desc)
        offer_template_doc = load_offer_letter_template(template)

        # Check similarity
        results = check_similarity(job_desc_text, resumes)

        # Initialize lists for the output
        analysis_results = ["Analysis Results:"]
        interview_messages = []
        offer_files = []

        eligible_candidates = [result for result in results if result[2] == "Eligible"]

        # Process each resume's similarity
        for idx, (filename, similarity, eligibility, extracted_name) in enumerate(results, start=1):
            candidate_label = f"Candidate {idx}"
            similarity_percentage = similarity * 100
            analysis_results.append(f"{candidate_label}, Similarity Percentage: {similarity_percentage:.2f}% - Status: {eligibility}")

            # If interview is scheduled and "Yes" is selected
            if interview_choice == "Yes" and eligibility == "Eligible" and extracted_name:
                is_valid, date_obj, time_obj = validate_date_time(interview_date, interview_time)
                if is_valid:
                    interview_msg = schedule_interview(candidate_label, interview_date, interview_time)
                    interview_messages.append(interview_msg)

                    # Ask the user if they want to generate the offer letter
                    if generate_offer_choice == "Yes":
                        offer_file = create_offer_letter(candidate_name, role, "XYZ Corp. Pakistan", joining_date, offer_template_doc)
                        offer_files.append(offer_file)
                    else:
                        interview_messages.append(f"Offer letter not generated for {candidate_label}.")
                else:
                    interview_messages.append(f"Invalid date or time format for {candidate_label}. Use YYYY-MM-DD for date and HH:MM AM/PM for time.")

        # Handling for the case when no candidates are eligible
        if not eligible_candidates:
            interview_messages.append("No eligible candidates for interview scheduling.")

        # Prepare interview schedule output
        if interview_messages:
            interview_messages.insert(0, "Interview Schedule:")
            interview_output = "\n".join(interview_messages)
        else:
            interview_output = "No interviews scheduled."

        # Prepare the offer letters output
        if offer_files:
            analysis_results.append("\nGenerated Offer Letters:")
            for idx, offer_file in enumerate(offer_files, start=1):
                analysis_results.append(f"- Candidate {idx} Offer Letter")

        # Join and return the results as formatted text
        analysis_output = "\n".join(analysis_results)
        return analysis_output, interview_output, offer_files

    except Exception as e:
        # Return any errors encountered during processing
        return f"Error processing files: {str(e)}", None
        # Gradio Interface Components
job_desc_input = gr.File(label="Upload Job Description (TXT)", type="filepath")
template_input = gr.File(label="Upload Offer Letter Template (DOCX)", type="filepath")
resumes_input = gr.Files(label="Upload Resumes (TXT, DOCX, PDF)", type="filepath")

interview_choice_input = gr.Radio(["Yes", "No"], label="Schedule Interview?")
interview_date_input = gr.Textbox(label="Interview Date (YYYY-MM-DD)", placeholder="Enter date in YYYY-MM-DD format")
interview_time_input = gr.Textbox(label="Interview Time (HH:MM AM/PM)", placeholder="Enter time in HH:MM AM/PM format")

generate_offer_choice_input = gr.Radio(["Yes", "No"], label="Generate Offer Letter?")
role_input = gr.Textbox(label="Enter Role")
joining_date_input = gr.Textbox(label="Enter Joining Date (YYYY-MM-DD)", placeholder="Enter joining date in YYYY-MM-DD format")
candidate_name_input = gr.Textbox(label="Enter Candidate Name", placeholder="Enter candidate's name")

# Gradio Outputs
results_output = gr.Markdown(label="Analysis Results")
interview_output = gr.Markdown(label="Interview Schedule")
offer_letters_output = gr.Files(label="Generated Offer Letters")

# Gradio Interface
interface = gr.Interface(
    fn=process_files,
    inputs=[job_desc_input, template_input, resumes_input, interview_choice_input, interview_date_input, interview_time_input, generate_offer_choice_input, role_input, joining_date_input, candidate_name_input],
    outputs=[results_output, interview_output, offer_letters_output],
    title="HR Assistant - Resume Screening & Interview Scheduling",
    description="Upload job description, template, and resumes to screen candidates, schedule interviews, and generate offer letters."
)

interface.launch()
